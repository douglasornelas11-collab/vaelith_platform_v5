from __future__ import annotations

import csv
import difflib
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import shutil
import sqlite3
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

import pdfplumber
import uvicorn
from docx import Document
from fastapi import Cookie, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BASE = Path(__file__).resolve().parent
DATA = BASE / "data"
UPLOADS = DATA / "uploads"
SAMPLES = BASE / "sample_files"
DB_PATH = DATA / "vaelith.db"
DATA.mkdir(exist_ok=True)
UPLOADS.mkdir(exist_ok=True)

APP_VERSION = "5.0-pilot"
SESSION_DAYS = int(os.getenv("SESSION_DAYS", "14"))
MAX_UPLOAD_MB = int(os.getenv("MAX_UPLOAD_MB", "750"))
COOKIE_SECURE = os.getenv("COOKIE_SECURE", "0") == "1"

app = FastAPI(title="VAELITH LABS — Plataforma de Compatibilização", version=APP_VERSION)

ACCEPTED = {
    ".ifc", ".rvt", ".dwg", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".xlsm",
    ".csv", ".mpp", ".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff", ".json"
}

DISCIPLINES = [
    "Arquitetura", "Estrutura", "Elétrica", "Hidráulica", "Sanitária", "Incêndio",
    "Climatização", "Interiores", "Fachada", "Telecomunicações", "Gás", "Paisagismo",
    "Orçamento", "Planejamento", "Escopo e memoriais", "Não identificada"
]

DEFAULT_EXPECTED_DISCIPLINES = [
    "Arquitetura", "Estrutura", "Elétrica", "Hidráulica", "Sanitária", "Incêndio",
    "Climatização", "Orçamento", "Planejamento", "Escopo e memoriais"
]

DISCIPLINE_KEYS = [
    ("ARQ", "Arquitetura"), ("ARQUIT", "Arquitetura"), ("ESTR", "Estrutura"),
    ("EST_", "Estrutura"), ("ELE", "Elétrica"), ("ELÉ", "Elétrica"),
    ("HID", "Hidráulica"), ("SANIT", "Sanitária"), ("ESG", "Sanitária"),
    ("INCEND", "Incêndio"), ("PCI", "Incêndio"), ("HVAC", "Climatização"),
    ("CLIMA", "Climatização"), ("AR COND", "Climatização"), ("ORC", "Orçamento"),
    ("ORÇ", "Orçamento"), ("CRONO", "Planejamento"), ("PLAN", "Planejamento"),
    ("INTERIOR", "Interiores"), ("FACH", "Fachada"), ("TELECOM", "Telecomunicações"),
    ("DADOS", "Telecomunicações"), ("GAS", "Gás"), ("GÁS", "Gás"),
    ("PAISAG", "Paisagismo"), ("MEMORIAL", "Escopo e memoriais"),
    ("ESCOPO", "Escopo e memoriais"), ("SPEC", "Escopo e memoriais")
]

ROOT_TYPES = {
    "IFCPROJECT", "IFCSITE", "IFCBUILDING", "IFCBUILDINGSTOREY", "IFCSPACE",
    "IFCWALL", "IFCWALLSTANDARDCASE", "IFCDOOR", "IFCWINDOW", "IFCSLAB", "IFCBEAM",
    "IFCCOLUMN", "IFCPIPESEGMENT", "IFCDUCTSEGMENT", "IFCCABLECARRIERSEGMENT",
    "IFCFLOWTERMINAL", "IFCFURNISHINGELEMENT", "IFCSTAIR", "IFCROOF", "IFCOPENINGELEMENT",
    "IFCPLATE", "IFCMEMBER", "IFCFOOTING", "IFCPILE", "IFCRAILING", "IFCCURTAINWALL",
    "IFCSANITARYTERMINAL", "IFCLIGHTFIXTURE", "IFCUNITARYEQUIPMENT", "IFCAIRTERMINAL",
    "IFCFLOWSEGMENT", "IFCFLOWFITTING", "IFCDISTRIBUTIONELEMENT", "IFCSPACEHEATER"
}

# Conceitos usados para cruzar projeto, orçamento, cronograma e escopo.
CONCEPTS = [
    {"key": "portas_vãos", "label": "Portas e vãos", "terms": ["porta", "portas", "vão", "vao", "abertura", "esquadria"], "disciplines": ["Arquitetura", "Estrutura"]},
    {"key": "estrutura", "label": "Estrutura de concreto/metálica", "terms": ["pilar", "viga", "laje", "concreto", "estrutura", "fundação", "fundacao", "metálica", "metalica"], "disciplines": ["Estrutura"]},
    {"key": "alvenaria", "label": "Alvenaria e fechamentos", "terms": ["alvenaria", "parede", "drywall", "fechamento", "divisória", "divisoria"], "disciplines": ["Arquitetura"]},
    {"key": "eletrica", "label": "Instalações elétricas", "terms": ["elétrica", "eletrica", "eletroduto", "eletrocalha", "tomada", "interruptor", "luminária", "luminaria", "quadro"], "disciplines": ["Elétrica"]},
    {"key": "hidraulica", "label": "Instalações hidráulicas", "terms": ["hidráulica", "hidraulica", "água", "agua", "tubulação", "tubulacao", "registro", "barrilete"], "disciplines": ["Hidráulica"]},
    {"key": "sanitaria", "label": "Esgoto e drenagem", "terms": ["esgoto", "sanitária", "sanitaria", "drenagem", "ralo", "caixa sifonada", "pluvial"], "disciplines": ["Sanitária"]},
    {"key": "hvac", "label": "Climatização e ventilação", "terms": ["climatização", "climatizacao", "duto", "evaporadora", "condensadora", "ventilação", "ventilacao", "ar condicionado"], "disciplines": ["Climatização", "Elétrica", "Sanitária"]},
    {"key": "incendio", "label": "Prevenção e combate a incêndio", "terms": ["sprinkler", "hidrante", "incêndio", "incendio", "detector", "alarme", "rota de fuga", "saída de emergência", "saida de emergencia"], "disciplines": ["Incêndio", "Arquitetura"]},
    {"key": "forro", "label": "Forros e interfaces superiores", "terms": ["forro", "sanca", "plenum", "tabica"], "disciplines": ["Arquitetura", "Elétrica", "Climatização", "Incêndio"]},
    {"key": "acabamentos", "label": "Acabamentos", "terms": ["pintura", "revestimento", "rodapé", "rodape", "porcelanato", "cerâmica", "ceramica"], "disciplines": ["Arquitetura", "Interiores"]},
    {"key": "telecom", "label": "Telecomunicações e dados", "terms": ["dados", "telecom", "cabeamento", "rack", "rede", "cftv"], "disciplines": ["Telecomunicações", "Elétrica"]},
    {"key": "fachada", "label": "Fachadas e esquadrias", "terms": ["fachada", "pele de vidro", "acm", "caixilho", "esquadria", "brise"], "disciplines": ["Fachada", "Arquitetura", "Estrutura"]},
    {"key": "equipamentos", "label": "Equipamentos", "terms": ["equipamento", "bomba", "gerador", "elevador", "transformador", "compressor"], "disciplines": ["Elétrica", "Hidráulica", "Climatização"]},
]


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def format_brl(value: float | int | None) -> str:
    if value is None:
        return "Não calculável"
    number = f"{float(value):,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    return f"R$ {number}"


def conn() -> sqlite3.Connection:
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    c.execute("PRAGMA foreign_keys=ON")
    c.execute("PRAGMA journal_mode=WAL")
    return c


def hash_password(password: str, salt: bytes | None = None) -> tuple[str, str]:
    salt = salt or secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 240_000)
    return salt.hex(), digest.hex()


def verify_password(password: str, salt_hex: str, digest_hex: str) -> bool:
    _, digest = hash_password(password, bytes.fromhex(salt_hex))
    return hmac.compare_digest(digest, digest_hex)


def init_db() -> None:
    with conn() as c:
        c.executescript("""
        CREATE TABLE IF NOT EXISTS users(
            id TEXT PRIMARY KEY, name TEXT NOT NULL, email TEXT UNIQUE NOT NULL,
            salt TEXT NOT NULL, password_hash TEXT NOT NULL, created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS sessions(
            token TEXT PRIMARY KEY, user_id TEXT NOT NULL, expires_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS projects(
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL, name TEXT NOT NULL,
            client TEXT NOT NULL DEFAULT '', location TEXT NOT NULL DEFAULT '', phase TEXT NOT NULL DEFAULT 'Pré-obra',
            description TEXT NOT NULL DEFAULT '', expected_disciplines_json TEXT NOT NULL DEFAULT '[]',
            baseline_status TEXT NOT NULL DEFAULT 'Não aprovada', baseline_analysis_id TEXT,
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS files(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL, ext TEXT NOT NULL, size_bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL, discipline TEXT NOT NULL, revision TEXT NOT NULL,
            category TEXT NOT NULL, status TEXT NOT NULL, summary TEXT NOT NULL,
            details_json TEXT NOT NULL, uploaded_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS changes(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, code TEXT NOT NULL,
            title TEXT NOT NULL, request_text TEXT NOT NULL, reason TEXT NOT NULL DEFAULT '',
            location TEXT NOT NULL, element TEXT NOT NULL, stage TEXT NOT NULL,
            base_deadline TEXT, status TEXT NOT NULL DEFAULT 'Em análise',
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS analyses(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, mode TEXT NOT NULL,
            change_id TEXT, result_json TEXT NOT NULL, created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS issues(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, analysis_id TEXT,
            code TEXT NOT NULL, source_type TEXT NOT NULL, category TEXT NOT NULL,
            severity TEXT NOT NULL, title TEXT NOT NULL, detail TEXT NOT NULL,
            location TEXT NOT NULL DEFAULT '', floor TEXT NOT NULL DEFAULT '',
            disciplines_json TEXT NOT NULL DEFAULT '[]', evidence_json TEXT NOT NULL DEFAULT '{}',
            confidence TEXT NOT NULL DEFAULT 'Estimado', status TEXT NOT NULL DEFAULT 'Aberta',
            responsible TEXT NOT NULL DEFAULT '', resolution TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        """)
        # Usuário de demonstração.
        user = c.execute("SELECT * FROM users WHERE email=?", ("demo@vaelithlabs.com.br",)).fetchone()
        if not user:
            uid = uuid4().hex
            salt, ph = hash_password("vaelith")
            c.execute("INSERT INTO users VALUES(?,?,?,?,?,?)", (uid, "Usuário Demo", "demo@vaelithlabs.com.br", salt, ph, now_iso()))
        else:
            uid = user["id"]
        project = c.execute("SELECT * FROM projects WHERE user_id=? AND name=?", (uid, "Empreendimento demonstrativo")).fetchone()
        if not project:
            pid = uuid4().hex
            c.execute("INSERT INTO projects VALUES(?,?,?,?,?,?,?,?,?,?,?,?)", (
                pid, uid, "Empreendimento demonstrativo", "Cliente exemplo", "Betim/MG", "Pré-obra",
                "Projeto piloto para validar compatibilização geral, revisões, orçamento, cronograma e mudanças.",
                json.dumps(DEFAULT_EXPECTED_DISCIPLINES, ensure_ascii=False), "Não aprovada", None, now_iso(), now_iso()
            ))
        else:
            pid = project["id"]
    seed_demo_files(pid)


def seed_demo_files(project_id: str) -> None:
    if not SAMPLES.exists():
        return
    with conn() as c:
        count = c.execute("SELECT COUNT(*) FROM files WHERE project_id=?", (project_id,)).fetchone()[0]
    if count:
        return
    mapping = {
        "ARQ_R00.ifc": ("Arquitetura", "R00"), "ARQ_R01.ifc": ("Arquitetura", "R01"),
        "EST_R01.ifc": ("Estrutura", "R01"), "ORC_R01.xlsx": ("Orçamento", "R01"),
        "CRONO_R01.xlsx": ("Planejamento", "R01"), "ARQ_MEMORIAL_R01.docx": ("Escopo e memoriais", "R01"),
        "ARQ_NOTA_R01.pdf": ("Escopo e memoriais", "R01")
    }
    for src in SAMPLES.iterdir():
        if not src.is_file():
            continue
        fid = uuid4().hex
        target = UPLOADS / f"{fid}{src.suffix.lower()}"
        shutil.copy2(src, target)
        raw = target.read_bytes()
        status, summary, details = analyze_file(target, src.name)
        discipline, revision = mapping.get(src.name, (infer_discipline(src.name), infer_revision(src.name)))
        category = infer_category(src.suffix.lower(), discipline)
        with conn() as c:
            c.execute("INSERT INTO files VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (
                fid, project_id, src.name, target.name, src.suffix.lower(), len(raw), hashlib.sha256(raw).hexdigest(),
                discipline, revision, category, status, summary, json.dumps(details, ensure_ascii=False), now_iso()
            ))



def require_user(session: str | None) -> sqlite3.Row:
    if not session:
        raise HTTPException(401, "Faça login para continuar")
    with conn() as c:
        row = c.execute("""
            SELECT u.* FROM sessions s JOIN users u ON u.id=s.user_id
            WHERE s.token=? AND s.expires_at>?
        """, (session, now_iso())).fetchone()
    if not row:
        raise HTTPException(401, "Sessão expirada")
    return row


def project_for_user(project_id: str, user_id: str) -> sqlite3.Row:
    with conn() as c:
        p = c.execute("SELECT * FROM projects WHERE id=? AND user_id=?", (project_id, user_id)).fetchone()
    if not p:
        raise HTTPException(404, "Empreendimento não encontrado")
    return p


def project_dict(row: sqlite3.Row) -> dict[str, Any]:
    d = dict(row)
    d["expectedDisciplines"] = json.loads(d.pop("expected_disciplines_json") or "[]")
    return d


def infer_revision(name: str) -> str:
    upper = name.upper()
    m = re.search(r"(?:^|[^A-Z0-9])(?:REV(?:IS[AÃ]O)?[ _.-]*|R)(\d{1,3})(?:[^A-Z0-9]|$)", upper)
    return f"R{int(m.group(1)):02d}" if m else "Não identificada"


def infer_discipline(name: str) -> str:
    upper = name.upper()
    for key, value in DISCIPLINE_KEYS:
        if key in upper:
            return value
    return "Não identificada"


def infer_category(ext: str, discipline: str) -> str:
    if ext in {".ifc", ".rvt"}:
        return "Modelo BIM"
    if ext in {".dwg", ".pdf", ".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff"}:
        return "Desenho ou evidência"
    if discipline == "Orçamento":
        return "Orçamento"
    if discipline == "Planejamento" or ext == ".mpp":
        return "Cronograma"
    if ext in {".doc", ".docx"}:
        return "Escopo e memorial"
    if ext in {".xls", ".xlsx", ".xlsm", ".csv"}:
        return "Planilha técnica"
    return "Documento"


def format_bytes(n: int) -> str:
    x = float(n)
    units = ["B", "KB", "MB", "GB"]
    i = 0
    while x >= 1024 and i < len(units) - 1:
        x /= 1024
        i += 1
    return f"{x:.1f} {units[i]}" if i else f"{int(x)} B"


def normalize_ifc_line(line: str) -> str:
    return re.sub(r"\s+", "", line).upper()


def parse_ifc(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    schema_m = re.search(r"FILE_SCHEMA\s*\(\s*\(\s*'([^']+)'", text, re.I)
    entity_counts = Counter(re.findall(r"=\s*(IFC[A-Z0-9_]+)\s*\(", text, re.I))
    elements: dict[str, dict[str, Any]] = {}
    duplicate_guids: list[str] = []
    storeys: list[str] = []
    spaces: list[str] = []
    for m in re.finditer(r"#(\d+)\s*=\s*(IFC[A-Z0-9_]+)\s*\((.*?)\);", text, re.I | re.S):
        express_id, entity, args = m.group(1), m.group(2).upper(), m.group(3)
        strings = re.findall(r"'((?:[^']|'')*)'", args)
        if entity == "IFCBUILDINGSTOREY" and strings:
            storeys.append(strings[2] if len(strings) > 2 else strings[-1])
        if entity == "IFCSPACE" and strings:
            spaces.append(strings[2] if len(strings) > 2 else strings[-1])
        if entity not in ROOT_TYPES:
            continue
        guid = strings[0] if strings else f"EXPRESS-{express_id}"
        name = strings[2] if len(strings) > 2 else (strings[1] if len(strings) > 1 else "")
        rec = {
            "guid": guid, "expressId": int(express_id), "entity": entity, "name": name,
            "signature": hashlib.sha1(normalize_ifc_line(m.group(0)).encode()).hexdigest(),
        }
        if guid in elements:
            duplicate_guids.append(guid)
        elements[guid] = rec
    return {
        "schema": schema_m.group(1) if schema_m else None,
        "entityCounts": entity_counts.most_common(60), "entityTotal": sum(entity_counts.values()),
        "elements": elements, "elementCount": len(elements), "duplicateGuids": duplicate_guids[:200],
        "storeys": sorted(set(x for x in storeys if x))[:100], "spaces": sorted(set(x for x in spaces if x))[:500]
    }


def find_header_map(rows: list[list[Any]]) -> tuple[int | None, dict[str, int]]:
    aliases = {
        "code": ["codigo", "código", "item", "id"],
        "description": ["descricao", "descrição", "servico", "serviço", "atividade", "nome"],
        "quantity": ["quantidade", "qtd", "quant."], "unit": ["unidade", "un"],
        "unit_price": ["preco unitario", "preço unitário", "valor unitario", "valor unitário", "pu"],
        "total": ["total", "valor total", "subtotal"], "start": ["inicio", "início", "data inicio", "data início"],
        "end": ["fim", "termino", "término", "data fim"], "duration": ["duracao", "duração", "dias", "prazo"],
        "predecessor": ["predecessora", "predecessor", "dependencia", "dependência"],
        "responsible": ["responsavel", "responsável", "equipe", "empresa"]
    }
    for idx, row in enumerate(rows[:30]):
        normalized = [re.sub(r"\s+", " ", str(v or "").strip().lower()) for v in row]
        mapping: dict[str, int] = {}
        for key, names in aliases.items():
            for col, val in enumerate(normalized):
                if any(alias == val or alias in val for alias in names):
                    mapping[key] = col
                    break
        if "description" in mapping and ("quantity" in mapping or "duration" in mapping or "start" in mapping or "total" in mapping):
            return idx, mapping
    return None, {}


def parse_workbook(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets, budgets, schedules = [], [], []
    for ws in wb.worksheets[:40]:
        rows = [[c for c in row] for row in ws.iter_rows(values_only=True)]
        header_idx, mapping = find_header_map(rows)
        sheets.append({"name": ws.title, "rows": ws.max_row, "columns": ws.max_column,
                       "preview": [[str(v or "")[:120] for v in row[:12]] for row in rows[:8]]})
        if header_idx is None:
            continue
        for row in rows[header_idx + 1:]:
            desc = str(row[mapping["description"]] or "").strip() if mapping["description"] < len(row) else ""
            if not desc:
                continue
            if "quantity" in mapping or "unit_price" in mapping or "total" in mapping:
                budgets.append({
                    "sheet": ws.title, "code": row[mapping["code"]] if "code" in mapping and mapping["code"] < len(row) else "",
                    "description": desc, "quantity": row[mapping["quantity"]] if "quantity" in mapping and mapping["quantity"] < len(row) else None,
                    "unit": row[mapping["unit"]] if "unit" in mapping and mapping["unit"] < len(row) else "",
                    "unitPrice": row[mapping["unit_price"]] if "unit_price" in mapping and mapping["unit_price"] < len(row) else None,
                    "total": row[mapping["total"]] if "total" in mapping and mapping["total"] < len(row) else None,
                })
            if "duration" in mapping or "start" in mapping:
                schedules.append({
                    "sheet": ws.title, "code": row[mapping["code"]] if "code" in mapping and mapping["code"] < len(row) else "",
                    "activity": desc, "duration": row[mapping["duration"]] if "duration" in mapping and mapping["duration"] < len(row) else None,
                    "start": str(row[mapping["start"]]) if "start" in mapping and mapping["start"] < len(row) and row[mapping["start"]] is not None else None,
                    "end": str(row[mapping["end"]]) if "end" in mapping and mapping["end"] < len(row) and row[mapping["end"]] is not None else None,
                    "predecessor": row[mapping["predecessor"]] if "predecessor" in mapping and mapping["predecessor"] < len(row) else None,
                    "responsible": row[mapping["responsible"]] if "responsible" in mapping and mapping["responsible"] < len(row) else None,
                })
    return {"sheets": sheets, "budgetRows": budgets[:20000], "scheduleRows": schedules[:20000]}


def parse_csv(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    sample = text[:8000]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|") if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    header_idx, mapping = find_header_map(rows)
    result: dict[str, Any] = {"preview": rows[:20], "budgetRows": [], "scheduleRows": []}
    if header_idx is None:
        return result
    for row in rows[header_idx + 1:]:
        if not row or mapping["description"] >= len(row):
            continue
        desc = row[mapping["description"]].strip()
        if not desc:
            continue
        if "quantity" in mapping or "unit_price" in mapping or "total" in mapping:
            result["budgetRows"].append({
                "description": desc, "quantity": row[mapping["quantity"]] if "quantity" in mapping and mapping["quantity"] < len(row) else None,
                "unit": row[mapping["unit"]] if "unit" in mapping and mapping["unit"] < len(row) else "",
                "unitPrice": row[mapping["unit_price"]] if "unit_price" in mapping and mapping["unit_price"] < len(row) else None,
                "total": row[mapping["total"]] if "total" in mapping and mapping["total"] < len(row) else None,
            })
        if "duration" in mapping or "start" in mapping:
            result["scheduleRows"].append({
                "activity": desc, "duration": row[mapping["duration"]] if "duration" in mapping and mapping["duration"] < len(row) else None,
                "start": row[mapping["start"]] if "start" in mapping and mapping["start"] < len(row) else None,
                "end": row[mapping["end"]] if "end" in mapping and mapping["end"] < len(row) else None,
                "predecessor": row[mapping["predecessor"]] if "predecessor" in mapping and mapping["predecessor"] < len(row) else None,
            })
    return result


def analyze_file(path: Path, original: str) -> tuple[str, str, dict[str, Any]]:
    ext = path.suffix.lower()
    try:
        if ext == ".ifc":
            d = parse_ifc(path)
            return "Processado", f"IFC {d.get('schema') or ''}: {d['elementCount']} elementos e {d['entityTotal']} entidades.", d
        if ext in {".xlsx", ".xlsm"}:
            d = parse_workbook(path)
            return "Processado", f"{len(d['sheets'])} abas; {len(d['budgetRows'])} linhas de orçamento e {len(d['scheduleRows'])} atividades detectadas.", d
        if ext == ".csv":
            d = parse_csv(path)
            return "Processado", f"CSV: {len(d['budgetRows'])} linhas de orçamento e {len(d['scheduleRows'])} atividades detectadas.", d
        if ext == ".pdf":
            parts = []
            with pdfplumber.open(path) as pdf:
                count = len(pdf.pages)
                for page in pdf.pages[:80]:
                    parts.append(page.extract_text() or "")
            text = "\n".join(parts)
            return "Processado", f"PDF com {count} páginas; {len(text)} caracteres extraídos.", {"pages": count, "text": text[:400000]}
        if ext == ".docx":
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = [[[cell.text for cell in row.cells] for row in table.rows[:300]] for table in doc.tables[:50]]
            return "Processado", f"Word com {len(paras)} parágrafos e {len(doc.tables)} tabelas.", {"paragraphs": paras[:10000], "tables": tables}
        if ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
            with Image.open(path) as im:
                return "Processado", f"Imagem {im.width} × {im.height}px.", {"width": im.width, "height": im.height, "mode": im.mode}
        if ext in {".rvt", ".dwg"}:
            return "Conversão necessária", "Arquivo aceito e catalogado. Para maquete e análise geométrica, envie também a exportação IFC.", {"recommended": "IFC", "nativeGeometry": False}
        if ext == ".mpp":
            return "Conversão necessária", "Arquivo aceito. Exporte também para XLSX ou CSV para leitura estruturada nesta versão.", {"recommended": "XLSX/CSV"}
        if ext in {".doc", ".xls"}:
            return "Conversão necessária", f"Arquivo legado aceito. Converta para {'.docx' if ext == '.doc' else '.xlsx'} para leitura automática.", {}
        return "Catalogado", "Arquivo salvo para referência.", {}
    except Exception as exc:
        return "Erro de processamento", f"O arquivo foi salvo, mas a leitura falhou: {type(exc).__name__}.", {"error": str(exc)[:3000]}


init_db()


def file_rows(project_id: str) -> list[dict[str, Any]]:
    with conn() as c:
        rows = c.execute("SELECT * FROM files WHERE project_id=? ORDER BY uploaded_at", (project_id,)).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        d["details"] = json.loads(d.pop("details_json"))
        d["size"] = format_bytes(d["size_bytes"])
        d["downloadUrl"] = f"/api/files/{d['id']}/raw"
        out.append(d)
    return out


def issue_rows(project_id: str, analysis_id: str | None = None) -> list[dict[str, Any]]:
    with conn() as c:
        if analysis_id:
            rows = c.execute("SELECT * FROM issues WHERE project_id=? AND analysis_id=? ORDER BY created_at", (project_id, analysis_id)).fetchall()
        else:
            rows = c.execute("SELECT * FROM issues WHERE project_id=? ORDER BY created_at DESC LIMIT 1000", (project_id,)).fetchall()
    out = []
    for row in rows:
        d = dict(row)
        d["disciplines"] = json.loads(d.pop("disciplines_json") or "[]")
        d["evidence"] = json.loads(d.pop("evidence_json") or "{}")
        out.append(d)
    return out


def revision_number(rev: str) -> int:
    m = re.search(r"(\d+)", rev or "")
    return int(m.group(1)) if m else -1


def compare_ifc_files(a: dict[str, Any], b: dict[str, Any]) -> dict[str, Any]:
    ea = a.get("details", {}).get("elements", {})
    eb = b.get("details", {}).get("elements", {})
    ga, gb = set(ea), set(eb)
    added = [eb[g] for g in sorted(gb - ga)]
    removed = [ea[g] for g in sorted(ga - gb)]
    modified = [{"before": ea[g], "after": eb[g]} for g in sorted(ga & gb) if ea[g].get("signature") != eb[g].get("signature")]
    return {
        "discipline": b["discipline"], "from": a["revision"], "to": b["revision"],
        "fileFrom": a["original_name"], "fileTo": b["original_name"],
        "added": added[:3000], "removed": removed[:3000], "modified": modified[:3000],
        "counts": {"added": len(added), "removed": len(removed), "modified": len(modified)}
    }


def safe_float(v: Any) -> float | None:
    if v is None or v == "":
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace("R$", "").replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def normalize_words(text: str) -> set[str]:
    stop = {"de", "da", "do", "e", "em", "para", "com", "um", "uma", "o", "a", "os", "as", "no", "na", "por", "que", "dos", "das"}
    return {w for w in re.findall(r"[a-záàâãéêíóôõúç0-9]+", text.lower()) if len(w) > 2 and w not in stop}


def corpus_from_files(files: list[dict[str, Any]]) -> dict[str, str]:
    design_parts, document_parts, budget_parts, schedule_parts = [], [], [], []
    for f in files:
        design_parts.append(f["original_name"])
        details = f.get("details", {})
        for el in details.get("elements", {}).values():
            design_parts.append(f"{el.get('entity','')} {el.get('name','')}")
        if details.get("text"):
            document_parts.append(details["text"])
        if details.get("paragraphs"):
            document_parts.extend(details["paragraphs"])
        for table in details.get("tables", []):
            for row in table:
                document_parts.append(" ".join(str(x) for x in row))
        for row in details.get("budgetRows", []):
            budget_parts.append(str(row.get("description", "")))
        for row in details.get("scheduleRows", []):
            schedule_parts.append(str(row.get("activity", "")))
    return {
        "design": normalize_text(" ".join(design_parts)),
        "documents": normalize_text(" ".join(document_parts)),
        "budget": normalize_text(" ".join(budget_parts)),
        "schedule": normalize_text(" ".join(schedule_parts)),
    }


def contains_any(text: str, terms: list[str]) -> bool:
    return any(normalize_text(term) in text for term in terms)


def budget_summary(files: list[dict[str, Any]]) -> dict[str, Any]:
    rows = []
    for f in files:
        rows.extend([{**r, "file": f["original_name"]} for r in f.get("details", {}).get("budgetRows", [])])
    total = 0.0
    calculable = 0
    missing = 0
    duplicates: dict[str, list[dict[str, Any]]] = defaultdict(list)
    enriched = []
    for row in rows:
        qty = safe_float(row.get("quantity"))
        unit_price = safe_float(row.get("unitPrice"))
        row_total = safe_float(row.get("total"))
        if row_total is None and qty is not None and unit_price is not None:
            row_total = qty * unit_price
        if row_total is not None:
            total += row_total
            calculable += 1
        else:
            missing += 1
        key = re.sub(r"[^a-z0-9áàâãéêíóôõúç]+", " ", str(row.get("description", "")).lower()).strip()
        if key:
            duplicates[key].append(row)
        enriched.append({**row, "calculatedTotal": row_total})
    duplicate_groups = [v for v in duplicates.values() if len(v) > 1]
    return {
        "status": "ready" if rows else "missing", "rowCount": len(rows), "calculableRows": calculable,
        "missingValueRows": missing, "total": total if calculable else None,
        "duplicateGroups": [{"description": g[0].get("description"), "count": len(g)} for g in duplicate_groups[:100]],
        "rows": enriched[:5000]
    }


def schedule_summary(files: list[dict[str, Any]]) -> dict[str, Any]:
    rows = []
    for f in files:
        rows.extend([{**r, "file": f["original_name"]} for r in f.get("details", {}).get("scheduleRows", [])])
    with_duration = 0
    without_predecessor = 0
    total_duration = 0.0
    for row in rows:
        duration = safe_float(row.get("duration"))
        if duration is not None:
            with_duration += 1
            total_duration += duration
        if not str(row.get("predecessor") or "").strip():
            without_predecessor += 1
    return {
        "status": "ready" if rows else "missing", "rowCount": len(rows), "withDuration": with_duration,
        "withoutPredecessor": without_predecessor, "sumDurations": total_duration if with_duration else None,
        "rows": rows[:5000]
    }


def new_issue(category: str, severity: str, title: str, detail: str, *, disciplines: list[str] | None = None,
              confidence: str = "Confirmado", location: str = "", floor: str = "", evidence: dict[str, Any] | None = None,
              source_type: str = "analysis") -> dict[str, Any]:
    return {
        "category": category, "severity": severity, "title": title, "detail": detail,
        "disciplines": disciplines or [], "confidence": confidence, "location": location,
        "floor": floor, "evidence": evidence or {}, "sourceType": source_type
    }


def persist_issues(project_id: str, analysis_id: str, issues: list[dict[str, Any]]) -> list[dict[str, Any]]:
    persisted = []
    with conn() as c:
        for idx, issue in enumerate(issues, start=1):
            iid = uuid4().hex
            code = f"INT-{idx:04d}"
            created = now_iso()
            c.execute("INSERT INTO issues VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (
                iid, project_id, analysis_id, code, issue.get("sourceType", "analysis"), issue["category"], issue["severity"],
                issue["title"], issue["detail"], issue.get("location", ""), issue.get("floor", ""),
                json.dumps(issue.get("disciplines", []), ensure_ascii=False), json.dumps(issue.get("evidence", {}), ensure_ascii=False),
                issue.get("confidence", "Estimado"), "Aberta", "", "", created, created
            ))
            persisted.append({"id": iid, "code": code, "status": "Aberta", "responsible": "", "resolution": "", **issue})
    return persisted


def run_compatibility_analysis(project_id: str) -> dict[str, Any]:
    files = file_rows(project_id)
    with conn() as c:
        project = c.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    expected = json.loads(project["expected_disciplines_json"] or "[]")
    issues: list[dict[str, Any]] = []
    by_disc: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for f in files:
        by_disc[f["discipline"]].append(f)
        if f["status"] == "Erro de processamento":
            issues.append(new_issue("Qualidade do arquivo", "critical", f"Falha ao processar {f['original_name']}", f["summary"], confidence="Confirmado", evidence={"fileId": f["id"]}))
        elif f["status"] == "Conversão necessária":
            issues.append(new_issue("Formato", "warning", f"Conversão necessária: {f['original_name']}", f["summary"], confidence="Confirmado", evidence={"fileId": f["id"]}))
        if f["revision"] == "Não identificada":
            issues.append(new_issue("Controle de revisões", "high", "Revisão não identificada", f"Informe a revisão do arquivo {f['original_name']}.", evidence={"fileId": f["id"]}))
        if f["discipline"] == "Não identificada":
            issues.append(new_issue("Classificação", "high", "Disciplina não identificada", f"Classifique o arquivo {f['original_name']} antes da liberação da base.", evidence={"fileId": f["id"]}))
        if f["ext"] == ".ifc":
            for guid in f.get("details", {}).get("duplicateGuids", []):
                issues.append(new_issue("Qualidade IFC", "high", "GUID duplicado no modelo", f"O arquivo {f['original_name']} possui o GUID {guid} repetido.", disciplines=[f["discipline"]], evidence={"fileId": f["id"], "guid": guid}))
            if f.get("details", {}).get("elementCount", 0) == 0:
                issues.append(new_issue("Qualidade IFC", "critical", "IFC sem elementos analisáveis", f"O arquivo {f['original_name']} não contém elementos reconhecidos para análise.", disciplines=[f["discipline"]], evidence={"fileId": f["id"]}))

    present_disc = {d for d in by_disc if d != "Não identificada"}
    for disc in expected:
        if disc not in present_disc:
            issues.append(new_issue("Documento ausente", "high", f"Disciplina não enviada: {disc}", "A compatibilização não consegue verificar integralmente esta disciplina.", disciplines=[disc], confidence="Confirmado"))

    latest = {d: max((revision_number(f["revision"]) for f in fs), default=-1) for d, fs in by_disc.items() if d != "Não identificada"}
    known_revs = [v for v in latest.values() if v >= 0]
    if known_revs:
        top = max(known_revs)
        for d, rev in latest.items():
            if rev >= 0 and rev < top and d not in {"Orçamento", "Planejamento", "Escopo e memoriais"}:
                issues.append(new_issue("Compatibilidade de versões", "high", f"{d} pode estar desatualizado", f"Última revisão identificada: R{rev:02d}; maior revisão técnica do empreendimento: R{top:02d}.", disciplines=[d], confidence="Provável"))

    duplicate_hashes: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for f in files:
        duplicate_hashes[f["sha256"]].append(f)
    for same in duplicate_hashes.values():
        if len(same) > 1:
            issues.append(new_issue("Duplicidade", "info", "Arquivos binariamente idênticos", ", ".join(x["original_name"] for x in same), evidence={"fileIds": [x["id"] for x in same]}))

    ifc_comparisons = []
    for disc, fs in by_disc.items():
        ifcs = [f for f in fs if f["ext"] == ".ifc" and f["status"] == "Processado"]
        ifcs.sort(key=lambda f: revision_number(f["revision"]))
        for a, b in zip(ifcs, ifcs[1:]):
            comp = compare_ifc_files(a, b)
            ifc_comparisons.append(comp)
            count = sum(comp["counts"].values())
            if count:
                issues.append(new_issue("Alteração entre revisões", "info", f"{disc}: {a['revision']} → {b['revision']}", f"{comp['counts']['added']} elementos adicionados, {comp['counts']['removed']} removidos e {comp['counts']['modified']} modificados.", disciplines=[disc], confidence="Calculado", evidence={"comparison": comp["counts"], "fileFrom": a["id"], "fileTo": b["id"]}))

    text_comparisons = []
    for disc, fs in by_disc.items():
        docs = [f for f in fs if f["ext"] in {".pdf", ".docx"} and f["status"] == "Processado"]
        docs.sort(key=lambda f: revision_number(f["revision"]))
        for a, b in zip(docs, docs[1:]):
            ta = a["details"].get("text", "\n".join(a["details"].get("paragraphs", [])))
            tb = b["details"].get("text", "\n".join(b["details"].get("paragraphs", [])))
            ratio = difflib.SequenceMatcher(None, ta[:200000], tb[:200000]).ratio() if ta or tb else 1.0
            text_comparisons.append({"discipline": disc, "from": a["revision"], "to": b["revision"], "similarity": round(ratio * 100, 1), "fileFrom": a["original_name"], "fileTo": b["original_name"]})

    budget = budget_summary(files)
    schedule = schedule_summary(files)
    if budget["status"] == "missing":
        issues.append(new_issue("Orçamento", "high", "Orçamento estruturado não localizado", "Envie XLSX ou CSV com descrição, quantidade, unidade, preço unitário ou total.", disciplines=["Orçamento"], confidence="Confirmado"))
    else:
        if budget["missingValueRows"]:
            issues.append(new_issue("Orçamento", "warning", "Itens sem valor calculável", f"{budget['missingValueRows']} de {budget['rowCount']} linhas não possuem total nem quantidade × preço unitário calculável.", disciplines=["Orçamento"], confidence="Calculado"))
        for dup in budget["duplicateGroups"][:20]:
            issues.append(new_issue("Orçamento", "warning", "Possível item duplicado", f"“{dup['description']}” aparece {dup['count']} vezes. Validar se são serviços distintos.", disciplines=["Orçamento"], confidence="Provável"))
    if schedule["status"] == "missing":
        issues.append(new_issue("Cronograma", "high", "Cronograma estruturado não localizado", "Envie XLSX ou CSV com atividade, duração, início/fim e predecessoras.", disciplines=["Planejamento"], confidence="Confirmado"))
    elif schedule["withoutPredecessor"] > max(3, int(schedule["rowCount"] * 0.6)):
        issues.append(new_issue("Cronograma", "warning", "Poucas dependências identificadas", f"{schedule['withoutPredecessor']} de {schedule['rowCount']} atividades não possuem predecessora reconhecida. O impacto no caminho crítico ficará limitado.", disciplines=["Planejamento"], confidence="Calculado"))

    corpus = corpus_from_files(files)
    scope_matrix = []
    for concept in CONCEPTS:
        in_design = contains_any(corpus["design"] + " " + corpus["documents"], concept["terms"])
        in_budget = contains_any(corpus["budget"], concept["terms"])
        in_schedule = contains_any(corpus["schedule"], concept["terms"])
        relevant = in_design or in_budget or in_schedule
        row = {"key": concept["key"], "label": concept["label"], "design": in_design, "budget": in_budget, "schedule": in_schedule, "disciplines": concept["disciplines"]}
        scope_matrix.append(row)
        if not relevant:
            continue
        if in_design and not in_budget and budget["status"] == "ready":
            issues.append(new_issue("Falha potencial de escopo", "high", f"{concept['label']} sem correspondência no orçamento", "O conceito aparece nos projetos ou memoriais, mas não foi localizado nos itens orçamentários. Validar nomenclaturas e eventual omissão.", disciplines=concept["disciplines"] + ["Orçamento"], confidence="Provável", evidence={"concept": concept["key"]}))
        if in_budget and not in_design:
            issues.append(new_issue("Divergência projeto × orçamento", "warning", f"{concept['label']} aparece no orçamento, mas não nos projetos lidos", "Pode haver item genérico, projeto ausente ou serviço sem representação nos arquivos enviados.", disciplines=concept["disciplines"] + ["Orçamento"], confidence="Possível", evidence={"concept": concept["key"]}))
        if (in_design or in_budget) and not in_schedule and schedule["status"] == "ready":
            issues.append(new_issue("Planejamento", "warning", f"{concept['label']} sem atividade correspondente", "Não foi localizada atividade equivalente no cronograma. Validar se o serviço está agregado em outro pacote ou foi omitido.", disciplines=concept["disciplines"] + ["Planejamento"], confidence="Provável", evidence={"concept": concept["key"]}))
        missing_related = [d for d in concept["disciplines"] if d not in present_disc]
        if in_design and missing_related:
            issues.append(new_issue("Coordenação multidisciplinar", "high", f"Disciplinas relacionadas ausentes em {concept['label']}", f"Foram identificadas referências a {concept['label'].lower()}, mas não foram enviados arquivos de: {', '.join(missing_related)}.", disciplines=concept["disciplines"], confidence="Provável", evidence={"concept": concept["key"]}))

    ifc_files = [f for f in files if f["ext"] == ".ifc" and f["status"] == "Processado"]
    ifc_disciplines = {f["discipline"] for f in ifc_files if f["discipline"] != "Não identificada"}
    geometric = {
        "status": "ready" if len(ifc_disciplines) >= 2 else "insufficient",
        "engine": "Browser envelope triage",
        "exactEngineAvailable": False,
        "reason": "Modelos IFC de ao menos duas disciplinas estão prontos para maquete federada e pré-clash no navegador." if len(ifc_disciplines) >= 2 else "Envie IFCs de pelo menos duas disciplinas para formar a maquete federada e executar a triagem geométrica.",
        "ifcFiles": [{"id": f["id"], "name": f["original_name"], "discipline": f["discipline"], "revision": f["revision"], "url": f"/api/files/{f['id']}/raw"} for f in ifc_files]
    }
    if geometric["status"] != "ready":
        issues.append(new_issue("Maquete 3D", "high", "Base geométrica insuficiente", geometric["reason"], confidence="Confirmado"))
    else:
        issues.append(new_issue("Maquete 3D", "info", "Executar pré-clash geométrico", "A compatibilização documental está concluída. Abra a maquete, execute o pré-clash e salve as ocorrências para vinculá-las ao relatório.", confidence="Confirmado"))

    score = 0
    score += 15 if files else 0
    score += 15 if all(f["discipline"] != "Não identificada" for f in files) and files else 0
    score += 10 if all(f["revision"] != "Não identificada" for f in files) and files else 0
    score += 20 if geometric["status"] == "ready" else 0
    score += 15 if budget["status"] == "ready" else 0
    score += 15 if schedule["status"] == "ready" else 0
    expected_present = len([d for d in expected if d in present_disc]) / max(1, len(expected))
    score += round(10 * expected_present)
    score = min(100, score)

    analysis_id = uuid4().hex
    persisted = persist_issues(project_id, analysis_id, issues)
    critical = len([i for i in persisted if i["severity"] == "critical"])
    high = len([i for i in persisted if i["severity"] == "high"])
    conclusion = (
        f"A VAELITH processou {len(files)} arquivos de {len(present_disc)} disciplinas e gerou {len(persisted)} ocorrências. "
        f"A base possui qualidade estimada de {score}%. Foram encontrados {critical} alerta(s) crítico(s) e {high} de alta prioridade. "
        + ("Os modelos IFC estão disponíveis para maquete federada e pré-clash geométrico. " if geometric["status"] == "ready" else "A maquete federada ainda não pode ser formada com segurança. ")
        + (f"O orçamento possui {budget['rowCount']} itens e valor-base reconhecido de {format_brl(budget['total'])}. " if budget["status"] == "ready" else "O orçamento estruturado ainda não foi reconhecido. ")
        + (f"O cronograma possui {schedule['rowCount']} atividades reconhecidas." if schedule["status"] == "ready" else "O cronograma estruturado ainda não foi reconhecido.")
    )
    result = {
        "id": analysis_id, "mode": "compatibility", "createdAt": now_iso(), "projectId": project_id,
        "files": files, "issues": persisted, "revisionMatrix": latest, "ifcComparisons": ifc_comparisons,
        "textComparisons": text_comparisons, "budget": budget, "schedule": schedule, "scopeMatrix": scope_matrix,
        "geometric": geometric, "dataQuality": score, "conclusion": conclusion,
        "metrics": {"files": len(files), "disciplines": len(present_disc), "issues": len(persisted), "critical": critical, "high": high,
                    "ifcModels": len(ifc_files), "budgetRows": budget["rowCount"], "scheduleRows": schedule["rowCount"]},
        "trace": [
            {"step": 1, "title": "Inventário", "detail": f"{len(files)} arquivos foram catalogados por formato, disciplina, categoria e revisão."},
            {"step": 2, "title": "Qualidade da base", "detail": "Foram verificados processamento, revisões, disciplinas, duplicidades e integridade básica dos IFCs."},
            {"step": 3, "title": "Revisões", "detail": f"{len(ifc_comparisons)} comparações IFC e {len(text_comparisons)} comparações documentais foram executadas."},
            {"step": 4, "title": "Escopo", "detail": "Projetos, memoriais, orçamento e cronograma foram cruzados por conceitos técnicos para localizar possíveis omissões e divergências."},
            {"step": 5, "title": "Orçamento", "detail": f"{budget['rowCount']} linhas foram reconhecidas; {budget['missingValueRows']} não possuem valor calculável."},
            {"step": 6, "title": "Cronograma", "detail": f"{schedule['rowCount']} atividades foram reconhecidas; {schedule['withoutPredecessor']} não possuem predecessora identificada."},
            {"step": 7, "title": "Geometria", "detail": geometric["reason"]},
        ],
        "limitations": [
            "As divergências semânticas são triagens e precisam de validação do coordenador de projetos.",
            "O pré-clash do navegador utiliza envelopes geométricos e pode conter falsos positivos.",
            "RVT e DWG são catalogados, mas a análise geométrica requer exportação IFC nesta versão.",
            "O valor do orçamento é a soma dos itens reconhecidos e não representa automaticamente custo de impacto."
        ]
    }
    with conn() as c:
        c.execute("INSERT INTO analyses VALUES(?,?,?,?,?,?)", (analysis_id, project_id, "compatibility", None, json.dumps(result, ensure_ascii=False), result["createdAt"]))
        c.execute("UPDATE projects SET updated_at=? WHERE id=?", (now_iso(), project_id))
    return result


def expand_change_words(words: set[str]) -> set[str]:
    expanded = set(words)
    groups = [
        ({"porta", "vão", "vao", "abertura"}, {"alvenaria", "demolição", "demolicao", "reboco", "pintura", "acabamento", "porta", "verga"}),
        ({"interruptor", "tomada", "elétrico", "eletrico", "elétrica", "eletrica", "eletroduto"}, {"instalações", "instalacoes", "elétrica", "eletrica", "remanejamento", "interruptor"}),
        ({"tubulação", "tubulacao", "hidráulica", "hidraulica", "esgoto"}, {"hidráulica", "hidraulica", "tubulação", "tubulacao", "instalações", "instalacoes", "drenagem"}),
        ({"parede", "alvenaria"}, {"parede", "alvenaria", "reboco", "pintura", "demolição", "demolicao"}),
        ({"forro", "climatização", "climatizacao", "duto"}, {"forro", "climatização", "climatizacao", "duto", "acabamento", "elétrica", "eletrica"}),
    ]
    for triggers, additions in groups:
        if words & triggers:
            expanded |= additions
    return expanded


def match_budget_change(change: sqlite3.Row, files: list[dict[str, Any]]) -> dict[str, Any]:
    budget_rows = []
    for f in files:
        budget_rows.extend([{**r, "file": f["original_name"]} for r in f.get("details", {}).get("budgetRows", [])])
    if not budget_rows:
        return {"status": "not_calculable", "reason": "Nenhuma planilha de orçamento estruturada foi vinculada.", "matches": [], "total": None}
    query = " ".join([change["title"], change["request_text"], change["reason"], change["element"], change["location"]])
    qwords = expand_change_words(normalize_words(query))
    scored = []
    for row in budget_rows:
        score = len(qwords & normalize_words(str(row.get("description", ""))))
        if not score:
            continue
        qty, pu, total = safe_float(row.get("quantity")), safe_float(row.get("unitPrice")), safe_float(row.get("total"))
        if total is None and qty is not None and pu is not None:
            total = qty * pu
        scored.append({**row, "matchScore": score, "calculatedTotal": total})
    scored.sort(key=lambda r: (-r["matchScore"], str(r.get("description"))))
    matches = scored[:60]
    values = [r["calculatedTotal"] for r in matches if r.get("calculatedTotal") is not None]
    return {"status": "calculated" if values else "partial", "reason": "Itens semelhantes foram localizados; a seleção e os quantitativos precisam de validação técnica." if matches else "Nenhum item correspondeu ao texto da mudança.", "matches": matches, "total": sum(values) if values else None}


def match_schedule_change(change: sqlite3.Row, files: list[dict[str, Any]]) -> dict[str, Any]:
    rows = []
    for f in files:
        rows.extend([{**r, "file": f["original_name"]} for r in f.get("details", {}).get("scheduleRows", [])])
    if not rows:
        return {"status": "not_calculable", "reason": "Nenhum cronograma estruturado foi vinculado.", "matches": [], "days": None}
    query = " ".join([change["title"], change["request_text"], change["reason"], change["element"], change["location"]])
    qwords = expand_change_words(normalize_words(query))
    scored = []
    for row in rows:
        score = len(qwords & normalize_words(str(row.get("activity", ""))))
        if score:
            scored.append({**row, "matchScore": score, "durationNumber": safe_float(row.get("duration"))})
    scored.sort(key=lambda r: (-r["matchScore"], str(r.get("activity"))))
    matches = scored[:60]
    durations = [r["durationNumber"] for r in matches if r.get("durationNumber") is not None]
    return {"status": "calculated" if durations else "partial", "reason": "Atividades relacionadas foram localizadas; o impacto líquido depende das predecessoras, folgas e estágio da obra." if matches else "Nenhuma atividade correspondeu ao texto da mudança.", "matches": matches, "days": max(durations) if durations else None}


def latest_analysis(project_id: str, mode: str) -> dict[str, Any] | None:
    with conn() as c:
        row = c.execute("SELECT id,result_json FROM analyses WHERE project_id=? AND mode=? ORDER BY created_at DESC LIMIT 1", (project_id, mode)).fetchone()
    if not row:
        return None
    result = json.loads(row["result_json"])
    result["issues"] = issue_rows(project_id, row["id"])
    return result


def run_change_analysis(project_id: str, change_id: str) -> dict[str, Any]:
    files = file_rows(project_id)
    with conn() as c:
        change = c.execute("SELECT * FROM changes WHERE id=? AND project_id=?", (change_id, project_id)).fetchone()
        project = c.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    if not change:
        raise HTTPException(404, "Mudança não encontrada")
    baseline = latest_analysis(project_id, "compatibility")
    issues: list[dict[str, Any]] = []
    if not baseline:
        issues.append(new_issue("Versão-base", "critical", "Compatibilização geral ainda não executada", "Execute e aprove a compatibilização geral antes de analisar impactos de mudanças.", confidence="Confirmado"))
    elif project["baseline_status"] == "Não aprovada":
        issues.append(new_issue("Versão-base", "high", "Versão-base não aprovada", "A análise usa a última compatibilização disponível, mas a base ainda não foi formalmente aprovada.", confidence="Confirmado"))

    budget = match_budget_change(change, files)
    schedule = match_schedule_change(change, files)
    query_words = expand_change_words(normalize_words(" ".join([change["title"], change["request_text"], change["reason"], change["element"], change["location"]])))
    impacted_disciplines = []
    for concept in CONCEPTS:
        if query_words & normalize_words(" ".join(concept["terms"])):
            impacted_disciplines.extend(concept["disciplines"])
    impacted_disciplines = sorted(set(impacted_disciplines))
    if not impacted_disciplines:
        impacted_disciplines = ["Arquitetura"]
        issues.append(new_issue("Escopo da mudança", "warning", "Disciplinas impactadas não identificadas com segurança", "A descrição da mudança precisa de mais elementos, localização e justificativa para ampliar a análise.", confidence="Possível"))
    else:
        issues.append(new_issue("Impacto multidisciplinar", "info", "Disciplinas potencialmente afetadas", ", ".join(impacted_disciplines), disciplines=impacted_disciplines, confidence="Provável"))

    if budget["total"] is None:
        issues.append(new_issue("Custo da mudança", "high", "Impacto financeiro não calculável", budget["reason"], disciplines=["Orçamento"], confidence="Confirmado" if budget["status"] == "not_calculable" else "Possível"))
    else:
        issues.append(new_issue("Custo da mudança", "warning", "Itens orçamentários relacionados", f"Foram relacionados itens que somam {format_brl(budget['total'])}. Esse valor não inclui automaticamente demolição, retrabalho, perdas, custos indiretos ou aceleração.", disciplines=["Orçamento"], confidence="Estimado"))
    if schedule["days"] is None:
        issues.append(new_issue("Prazo da mudança", "high", "Impacto no prazo não calculável", schedule["reason"], disciplines=["Planejamento"], confidence="Confirmado" if schedule["status"] == "not_calculable" else "Possível"))
    else:
        issues.append(new_issue("Prazo da mudança", "warning", "Atividades relacionadas localizadas", f"A maior duração localizada é de {schedule['days']:g} dia(s). O impacto líquido só pode ser fechado após análise de predecessoras, folgas e situação executada.", disciplines=["Planejamento"], confidence="Estimado"))

    analysis_id = uuid4().hex
    persisted = persist_issues(project_id, analysis_id, issues)
    conclusion = (
        f"A mudança {change['code']} foi analisada sobre a última base disponível. "
        f"Foram apontadas {len(impacted_disciplines)} disciplina(s) potencialmente afetadas: {', '.join(impacted_disciplines)}. "
        + (f"Os itens orçamentários relacionados somam preliminarmente {format_brl(budget['total'])}, sem representar ainda o impacto total. " if budget["total"] is not None else "O custo não pôde ser calculado com fidelidade. ")
        + (f"Foram localizadas atividades com até {schedule['days']:g} dia(s) de duração; o impacto líquido depende da lógica do cronograma." if schedule["days"] is not None else "O prazo não pôde ser recalculado com fidelidade.")
    )
    result = {
        "id": analysis_id, "mode": "change", "createdAt": now_iso(), "projectId": project_id,
        "change": dict(change), "baseline": {"status": project["baseline_status"], "analysisId": project["baseline_analysis_id"]},
        "issues": persisted, "impactedDisciplines": impacted_disciplines, "budget": budget, "schedule": schedule,
        "conclusion": conclusion, "dataQuality": baseline.get("dataQuality", 0) if baseline else 0,
        "trace": [
            {"step": 1, "title": "Referência", "detail": "A última compatibilização geral foi utilizada como base de comparação." if baseline else "Não existe compatibilização geral disponível."},
            {"step": 2, "title": "Solicitação", "detail": f"{change['title']} — {change['request_text']}"},
            {"step": 3, "title": "Disciplinas", "detail": ", ".join(impacted_disciplines)},
            {"step": 4, "title": "Orçamento", "detail": budget["reason"]},
            {"step": 5, "title": "Cronograma", "detail": schedule["reason"]},
            {"step": 6, "title": "Geometria", "detail": "O cenário 3D deve ser criado ou carregado para localizar espacialmente os impactos da mudança."},
        ],
        "limitations": ["A mudança não é aplicada automaticamente ao modelo IFC nesta versão piloto.", "Custos e prazo exigem validação profissional e dados estruturados."]
    }
    with conn() as c:
        c.execute("INSERT INTO analyses VALUES(?,?,?,?,?,?)", (analysis_id, project_id, "change", change_id, json.dumps(result, ensure_ascii=False), result["createdAt"]))
        c.execute("UPDATE changes SET status='Analisada', updated_at=? WHERE id=?", (now_iso(), change_id))
    return result


def analysis_for_export(project_id: str, mode: str) -> dict[str, Any]:
    analysis = latest_analysis(project_id, mode)
    if not analysis:
        analysis = run_compatibility_analysis(project_id) if mode == "compatibility" else None
    if not analysis:
        raise HTTPException(404, "Nenhuma análise disponível")
    return analysis


@app.get("/api/health")
def health():
    return {"ok": True, "version": APP_VERSION, "maxUploadMb": MAX_UPLOAD_MB}


@app.post("/api/auth/register")
def register(payload: dict[str, Any]):
    name = str(payload.get("name", "")).strip()
    email = str(payload.get("email", "")).strip().lower()
    password = str(payload.get("password", ""))
    if len(name) < 2 or "@" not in email or len(password) < 8:
        raise HTTPException(400, "Informe nome, e-mail válido e senha com ao menos 8 caracteres")
    uid = uuid4().hex
    salt, ph = hash_password(password)
    try:
        with conn() as c:
            c.execute("INSERT INTO users VALUES(?,?,?,?,?,?)", (uid, name, email, salt, ph, now_iso()))
    except sqlite3.IntegrityError:
        raise HTTPException(409, "Este e-mail já está cadastrado")
    return {"ok": True}


@app.post("/api/auth/login")
def login(payload: dict[str, Any]):
    email = str(payload.get("email", "")).strip().lower()
    password = str(payload.get("password", ""))
    with conn() as c:
        user = c.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
    if not user or not verify_password(password, user["salt"], user["password_hash"]):
        raise HTTPException(401, "E-mail ou senha incorretos")
    token = secrets.token_urlsafe(42)
    expires = (datetime.now(timezone.utc) + timedelta(days=SESSION_DAYS)).isoformat()
    with conn() as c:
        c.execute("DELETE FROM sessions WHERE expires_at<=?", (now_iso(),))
        c.execute("INSERT INTO sessions VALUES(?,?,?)", (token, user["id"], expires))
    response = Response(status_code=204)
    response.set_cookie("vaelith_session", token, httponly=True, samesite="lax", secure=COOKIE_SECURE, max_age=SESSION_DAYS * 86400)
    return response


@app.post("/api/auth/logout")
def logout(vaelith_session: str | None = Cookie(default=None)):
    if vaelith_session:
        with conn() as c:
            c.execute("DELETE FROM sessions WHERE token=?", (vaelith_session,))
    response = Response(status_code=204)
    response.delete_cookie("vaelith_session")
    return response


@app.get("/api/auth/me")
def me(vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    return {"id": u["id"], "name": u["name"], "email": u["email"]}


@app.get("/api/projects")
def get_projects(vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        rows = c.execute("SELECT * FROM projects WHERE user_id=? ORDER BY updated_at DESC", (u["id"],)).fetchall()
    return [project_dict(r) for r in rows]


@app.post("/api/projects")
def create_project(payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    name = str(payload.get("name", "")).strip()
    if not name:
        raise HTTPException(400, "Informe o nome do empreendimento")
    pid = uuid4().hex
    expected = payload.get("expectedDisciplines") or DEFAULT_EXPECTED_DISCIPLINES
    with conn() as c:
        c.execute("INSERT INTO projects VALUES(?,?,?,?,?,?,?,?,?,?,?,?)", (
            pid, u["id"], name, str(payload.get("client", "")), str(payload.get("location", "")),
            str(payload.get("phase", "Pré-obra")), str(payload.get("description", "")),
            json.dumps(expected, ensure_ascii=False), "Não aprovada", None, now_iso(), now_iso()
        ))
    return {"id": pid}


@app.patch("/api/projects/{project_id}")
def update_project(project_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    expected = payload.get("expectedDisciplines", json.loads(p["expected_disciplines_json"] or "[]"))
    with conn() as c:
        c.execute("""UPDATE projects SET name=?,client=?,location=?,phase=?,description=?,expected_disciplines_json=?,updated_at=? WHERE id=?""", (
            str(payload.get("name", p["name"])), str(payload.get("client", p["client"])),
            str(payload.get("location", p["location"])), str(payload.get("phase", p["phase"])),
            str(payload.get("description", p["description"])), json.dumps(expected, ensure_ascii=False), now_iso(), project_id
        ))
    return {"ok": True}


@app.get("/api/projects/{project_id}/state")
def project_state(project_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    with conn() as c:
        changes = [dict(x) for x in c.execute("SELECT * FROM changes WHERE project_id=? ORDER BY created_at DESC", (project_id,)).fetchall()]
    return {
        "project": project_dict(p), "files": file_rows(project_id), "changes": changes,
        "compatibility": latest_analysis(project_id, "compatibility"), "changeAnalysis": latest_analysis(project_id, "change"),
        "issues": issue_rows(project_id)
    }


@app.post("/api/projects/{project_id}/files")
async def upload_files(project_id: str, files: list[UploadFile] = File(...), discipline: str = Form(default=""),
                       revision: str = Form(default=""), category: str = Form(default=""),
                       vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    added = []
    for upload in files:
        original = Path(upload.filename or "arquivo").name
        ext = Path(original).suffix.lower()
        if ext not in ACCEPTED:
            raise HTTPException(415, f"Formato não aceito: {ext or 'sem extensão'}")
        fid = uuid4().hex
        target = UPLOADS / f"{fid}{ext}"
        sha = hashlib.sha256()
        size = 0
        with target.open("wb") as out:
            while chunk := await upload.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_UPLOAD_MB * 1024 * 1024:
                    out.close(); target.unlink(missing_ok=True)
                    raise HTTPException(413, f"Arquivo acima do limite de {MAX_UPLOAD_MB} MB")
                sha.update(chunk); out.write(chunk)
        status, summary, details = analyze_file(target, original)
        disc = discipline.strip() or infer_discipline(original)
        rev = revision.strip().upper() or infer_revision(original)
        cat = category.strip() or infer_category(ext, disc)
        with conn() as c:
            c.execute("INSERT INTO files VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (
                fid, project_id, original, target.name, ext, size, sha.hexdigest(), disc, rev, cat,
                status, summary, json.dumps(details, ensure_ascii=False), now_iso()
            ))
            c.execute("UPDATE projects SET updated_at=? WHERE id=?", (now_iso(), project_id))
        added.append({"id": fid, "name": original, "status": status, "summary": summary, "discipline": disc, "revision": rev, "category": cat})
    return {"ok": True, "added": added, "files": file_rows(project_id)}


@app.patch("/api/files/{file_id}")
def update_file(file_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
        if not row or row["user_id"] != u["id"]:
            raise HTTPException(404, "Arquivo não encontrado")
        c.execute("UPDATE files SET discipline=?,revision=?,category=? WHERE id=?", (
            str(payload.get("discipline", row["discipline"])), str(payload.get("revision", row["revision"])).upper(),
            str(payload.get("category", row["category"])), file_id
        ))
    return {"ok": True}


@app.delete("/api/files/{file_id}")
def delete_file(file_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
        if not row or row["user_id"] != u["id"]:
            raise HTTPException(404, "Arquivo não encontrado")
        (UPLOADS / row["stored_name"]).unlink(missing_ok=True)
        c.execute("DELETE FROM files WHERE id=?", (file_id,))
    return {"ok": True}


@app.get("/api/files/{file_id}/raw")
def raw_file(file_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
    if not row or row["user_id"] != u["id"]:
        raise HTTPException(404, "Arquivo não encontrado")
    return FileResponse(UPLOADS / row["stored_name"], filename=row["original_name"])


@app.post("/api/projects/{project_id}/compatibility")
def analyze_compatibility(project_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    return run_compatibility_analysis(project_id)


@app.post("/api/projects/{project_id}/baseline")
def approve_baseline(project_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    analysis = latest_analysis(project_id, "compatibility")
    if not analysis:
        raise HTTPException(400, "Execute a compatibilização geral antes de aprovar a versão-base")
    status = str(payload.get("status", "Aprovada com ressalvas"))
    if status not in {"Aprovada", "Aprovada com ressalvas", "Não aprovada"}:
        raise HTTPException(400, "Status inválido")
    if status == "Aprovada" and any(i["severity"] == "critical" and i.get("status") != "Resolvida" for i in analysis.get("issues", [])):
        raise HTTPException(409, "Existem alertas críticos; use aprovação com ressalvas ou resolva as ocorrências")
    with conn() as c:
        c.execute("UPDATE projects SET baseline_status=?,baseline_analysis_id=?,updated_at=? WHERE id=?", (status, analysis["id"], now_iso(), project_id))
    return {"ok": True, "status": status, "analysisId": analysis["id"]}


@app.post("/api/projects/{project_id}/changes")
def create_change(project_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    cid = uuid4().hex
    code = str(payload.get("code") or f"SM-{datetime.now().strftime('%Y%m%d-%H%M')}")
    with conn() as c:
        c.execute("INSERT INTO changes VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)", (
            cid, project_id, code, str(payload.get("title", "Mudança sem título")), str(payload.get("requestText", "")),
            str(payload.get("reason", "")), str(payload.get("location", "")), str(payload.get("element", "")),
            str(payload.get("stage", "Não informado")), str(payload.get("baseDeadline", "")) or None,
            "Em análise", now_iso(), now_iso()
        ))
    return {"id": cid, "code": code}


@app.post("/api/projects/{project_id}/changes/{change_id}/analyze")
def analyze_change(project_id: str, change_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    return run_change_analysis(project_id, change_id)


@app.post("/api/projects/{project_id}/clashes")
def save_clashes(project_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    clashes = payload.get("clashes") or []
    saved = []
    with conn() as c:
        for clash in clashes[:500]:
            iid = uuid4().hex
            code = f"3D-{uuid4().hex[:6].upper()}"
            disc = sorted(set([str(clash.get("disciplineA", "")), str(clash.get("disciplineB", ""))]) - {""})
            detail = f"Interseção preliminar entre {clash.get('elementA','elemento A')} e {clash.get('elementB','elemento B')}. Volume de envelope: {clash.get('volume',0):.5f} m³."
            created = now_iso()
            evidence = {"modelA": clash.get("modelA"), "modelB": clash.get("modelB"), "center": clash.get("center"), "volume": clash.get("volume")}
            c.execute("INSERT INTO issues VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (
                iid, project_id, None, code, "viewer", "Interferência geométrica", "high", "Pré-clash 3D", detail,
                str(clash.get("location", "")), str(clash.get("floor", "")), json.dumps(disc, ensure_ascii=False),
                json.dumps(evidence, ensure_ascii=False), "Estimado", "Aberta", "", "", created, created
            ))
            saved.append({"id": iid, "code": code})
    return {"ok": True, "saved": saved}


@app.patch("/api/issues/{issue_id}")
def update_issue(issue_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT i.*,p.user_id FROM issues i JOIN projects p ON p.id=i.project_id WHERE i.id=?", (issue_id,)).fetchone()
        if not row or row["user_id"] != u["id"]:
            raise HTTPException(404, "Ocorrência não encontrada")
        c.execute("UPDATE issues SET status=?,responsible=?,resolution=?,updated_at=? WHERE id=?", (
            str(payload.get("status", row["status"])), str(payload.get("responsible", row["responsible"])),
            str(payload.get("resolution", row["resolution"])), now_iso(), issue_id
        ))
    return {"ok": True}


def export_analysis(project: sqlite3.Row, analysis: dict[str, Any], fmt: str) -> Response:
    title = f"VAELITH — {project['name']}"
    issues = analysis.get("issues", [])
    if fmt == "json":
        return Response(json.dumps(analysis, ensure_ascii=False, indent=2), media_type="application/json", headers={"Content-Disposition": "attachment; filename=vaelith-relatorio.json"})
    if fmt == "xlsx":
        wb = Workbook(); ws = wb.active; ws.title = "Resumo"
        for row in [["VAELITH LABS", "Soluções em Engenharia"], ["Empreendimento", project["name"]], ["Tipo", "Compatibilização geral" if analysis["mode"] == "compatibility" else "Impacto de mudança"], ["Qualidade da base (%)", analysis.get("dataQuality")], ["Conclusão", analysis.get("conclusion")]]:
            ws.append(row)
        ws.column_dimensions["A"].width = 32; ws.column_dimensions["B"].width = 120
        ws["A1"].fill = ws["B1"].fill = PatternFill("solid", fgColor="CFFF47"); ws["A1"].font = Font(bold=True)
        sh = wb.create_sheet("Interferências")
        sh.append(["Código", "Severidade", "Categoria", "Título", "Detalhe", "Disciplinas", "Confiança", "Status", "Responsável"])
        for i in issues:
            sh.append([i.get("code"), i.get("severity"), i.get("category"), i.get("title"), i.get("detail"), ", ".join(i.get("disciplines", [])), i.get("confidence"), i.get("status"), i.get("responsible")])
        for cell in sh[1]: cell.fill = PatternFill("solid", fgColor="CFFF47"); cell.font = Font(bold=True)
        for row in sh.iter_rows():
            for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in sh.columns: sh.column_dimensions[col[0].column_letter].width = min(70, max(12, max(len(str(c.value or "")) for c in col) + 2))
        if analysis.get("budget", {}).get("rows"):
            b = wb.create_sheet("Orçamento")
            b.append(["Arquivo", "Código", "Descrição", "Quantidade", "Unidade", "Preço unitário", "Total"])
            for r in analysis["budget"]["rows"]:
                b.append([r.get("file"), r.get("code"), r.get("description"), r.get("quantity"), r.get("unit"), r.get("unitPrice"), r.get("calculatedTotal")])
        if analysis.get("schedule", {}).get("rows"):
            s = wb.create_sheet("Cronograma")
            s.append(["Arquivo", "Código", "Atividade", "Duração", "Início", "Fim", "Predecessora", "Responsável"])
            for r in analysis["schedule"]["rows"]:
                s.append([r.get("file"), r.get("code"), r.get("activity"), r.get("duration"), r.get("start"), r.get("end"), r.get("predecessor"), r.get("responsible")])
        out = io.BytesIO(); wb.save(out)
        return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=vaelith-relatorio.xlsx"})
    if fmt == "docx":
        doc = Document(); doc.add_heading("VAELITH LABS", 0); doc.add_paragraph("Soluções em Engenharia")
        doc.add_heading(project["name"], 1); doc.add_heading("Conclusão", 2); doc.add_paragraph(analysis.get("conclusion", ""))
        doc.add_heading("Linha de raciocínio", 2)
        for step in analysis.get("trace", []): doc.add_paragraph(f"{step['step']}. {step['title']}: {step['detail']}", style="List Number")
        doc.add_heading("Interferências e pendências", 2)
        for i in issues: doc.add_paragraph(f"{i.get('code','')} [{i.get('severity','')}] {i.get('title','')}: {i.get('detail','')}", style="List Bullet")
        out = io.BytesIO(); doc.save(out)
        return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=vaelith-relatorio.docx"})
    if fmt == "pdf":
        out = io.BytesIO(); styles = getSampleStyleSheet()
        story = [Paragraph(title, styles["Title"]), Paragraph("Soluções em Engenharia", styles["Heading2"]), Spacer(1, 6 * mm), Paragraph("Conclusão", styles["Heading2"]), Paragraph(analysis.get("conclusion", ""), styles["BodyText"]), Spacer(1, 4 * mm), Paragraph("Linha de raciocínio", styles["Heading2"])]
        for step in analysis.get("trace", []): story += [Paragraph(f"<b>{step['step']}. {step['title']}:</b> {step['detail']}", styles["BodyText"]), Spacer(1, 2 * mm)]
        story += [Spacer(1, 3 * mm), Paragraph("Interferências e pendências", styles["Heading2"])]
        data = [["Código", "Prioridade", "Título", "Detalhe"]] + [[i.get("code", ""), i.get("severity", ""), i.get("title", ""), i.get("detail", "")] for i in issues]
        table = Table(data, repeatRows=1, colWidths=[22 * mm, 25 * mm, 48 * mm, 85 * mm])
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#CFFF47")), ("GRID", (0, 0), (-1, -1), .3, colors.grey), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
        story.append(table)
        SimpleDocTemplate(out, pagesize=A4, leftMargin=15 * mm, rightMargin=15 * mm, topMargin=15 * mm, bottomMargin=15 * mm).build(story)
        return Response(out.getvalue(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=vaelith-relatorio.pdf"})
    raise HTTPException(400, "Formato não suportado")


@app.get("/api/projects/{project_id}/export/{mode}/{fmt}")
def export(project_id: str, mode: str, fmt: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    if mode not in {"compatibility", "change"}:
        raise HTTPException(400, "Modo inválido")
    return export_analysis(p, analysis_for_export(project_id, mode), fmt.lower())


@app.get("/")
def landing():
    return FileResponse(BASE / "index.html")


@app.get("/login")
def login_page():
    return FileResponse(BASE / "login.html")


@app.get("/app")
def app_page():
    return FileResponse(BASE / "app.html")


app.mount("/assets", StaticFiles(directory=BASE / "assets"), name="assets")

if __name__ == "__main__":
    uvicorn.run("server:app", host="0.0.0.0", port=int(os.getenv("PORT", "8080")), reload=False)
